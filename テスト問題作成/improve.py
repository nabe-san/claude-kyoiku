"""
テスト問題改善スクリプト

input/improve/ に置いたWord・PDFの問題ファイルを読み込み、
問題ごとに診断・改善案・代替案を生成してWord出力する。
教科書分析は既存のキャッシュ（教科書/分析キャッシュ/）を流用する。

使い方:
  1. input/improve/ に改善したい問題ファイル（.docx または .pdf）を置く
  2. python improve.py を実行
  3. output/reports/ に改善レポート Word が出力される
"""

import os, sys, re, unicodedata
from pathlib import Path
from datetime import datetime
import anthropic
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv

from generate import sprint, read_text_file, load_config, get_or_create_analysis, find_path

load_dotenv()

BASE_DIR    = Path(__file__).parent
IMPROVE_DIR = BASE_DIR / "input" / "improve"
OUTPUT_DIR  = BASE_DIR / "output" / "reports"
MODEL       = "claude-sonnet-4-6"


# ================================================================
# 問題のパース
# ================================================================

def parse_questions(text: str) -> list[str]:
    """テキストから問題を1問ずつ抽出する"""
    # 「問1」「問１」「**問1**」などのパターンで分割
    pattern = r'(?:^|\n)(?=\*{0,2}問\s*[0-9０-９]+)'
    parts = re.split(pattern, text, flags=re.MULTILINE)
    return [p.strip() for p in parts
            if p.strip() and re.match(r'\*{0,2}問\s*[0-9０-９]+', p.strip())]


# ================================================================
# 改善プロンプト
# ================================================================

CRITERIA = """\
## テスト問題の評価基準

### 問題形式
- 4択問題・年代整序問題・資料読み取り問題のみ使用
- 一問一答形式（「〜は誰か」「〜はいつか」）は禁止

### 知識の種類
- 概念的知識（歴史の流れ・背景・因果関係・意義）80%以上
- 事実的知識（年号・人名・単純暗記）20%以下

### 難易度
- 易しい問題は作らない（標準20%・難80%が目安）

### 選択肢の作り方
- 長さは30字程度を標準とする
- 歴史の流れ・背景・文脈の理解がないと正答できない内容にする
- 「明らかに間違い」な選択肢は避け、どれもそれらしく見える選択肢にする
- 1つだけ「惜しいが違う」選択肢を必ず入れる

### 問いの質
- 「なぜ〜か」「〜の結果どうなったか」「〜の意義は何か」を問う
- 複数の出来事・時代を比較・関連づける問いを積極的に入れる

### 文体
- 試験問題として自然な文体にする
- 「教科書分析テキストによれば」などの不自然な表現は禁止
"""


def build_improve_prompt(question: str, textbook_analysis: str) -> str:
    context = f"\n\n## 教科書分析テキスト（参考）\n{textbook_analysis}" if textbook_analysis else ""
    return f"""\
あなたは高校の歴史・公共の定期試験問題を改善するアシスタントです。
以下の評価基準をもとに、与えられた問題を分析し、改善案と代替案を提示してください。

{CRITERIA}{context}

---

## 改善対象の問題

{question}

---

## 出力形式（必ず以下の見出しを使うこと）

### 診断
この問題の良い点と改善すべき点を2〜4文で述べる。評価基準のどの項目に照らして改善が必要かを具体的に指摘する。

### 改善案
元の問題の意図を活かしつつ評価基準に沿って改善した問題を出力する。
問題文・選択肢ア〜エ・正答・解説をセットで出力する。

### 代替案
同じ単元・テーマを扱いつつ、異なる切り口で出題した代替問題を1問出力する。
問題文・選択肢ア〜エ・正答・解説をセットで出力する。
"""


# ================================================================
# Claude API 呼び出し
# ================================================================

def improve_question(question: str, textbook_analysis: str, idx: int) -> str:
    client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
    sprint(f"  [問{idx}] 改善案を生成中...")
    prompt = build_improve_prompt(question, textbook_analysis)
    response = client.messages.create(
        model=MODEL,
        max_tokens=2048,
        messages=[{"role": "user", "content": prompt}],
    )
    return response.content[0].text


# ================================================================
# Word 出力
# ================================================================

def _set_page_margins(doc: Document):
    for section in doc.sections:
        section.page_width    = Cm(21.0)
        section.page_height   = Cm(29.7)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)


def _set_line_spacing(para, multiple: float = 1.15):
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = multiple


def _set_font(run, name: str, size_pt: float, bold: bool = False, color: str = None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def _set_para_shading(para, fill_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    existing = pPr.find(qn("w:shd"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(shd)


def _add_section_label(doc: Document, label: str, fill_color: str):
    """色付きセクションラベル行"""
    p = doc.add_paragraph()
    _set_line_spacing(p)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0)
    _set_para_shading(p, fill_color)
    run = p.add_run(f"  {label}")
    _set_font(run, "游ゴシック", 9.5, bold=True, color="333333")


def _add_content_lines(doc: Document, text: str, fill_color: str):
    """本文を行ごとに段落として追加する"""
    for line in text.strip().splitlines():
        clean = line.replace("**", "").strip()
        if not clean:
            continue
        p = doc.add_paragraph()
        _set_line_spacing(p)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        _set_para_shading(p, fill_color)
        is_choice = clean.startswith(("ア．", "イ．", "ウ．", "エ．",
                                       "ア.", "イ.", "ウ.", "エ."))
        p.paragraph_format.left_indent = Cm(1.5) if is_choice else Cm(0.5)
        run = p.add_run(clean)
        _set_font(run, "游明朝", 10.5)


def _add_separator(doc: Document):
    """問題ブロック間の区切り線"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "thick")
    bottom.set(qn("w:sz"),    "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "595959")
    pb.append(bottom)
    pPr.append(pb)
    _set_line_spacing(p)


def _parse_sections(improvement: str) -> dict[str, str]:
    """Claude の出力を診断・改善案・代替案に分割する"""
    sections = {"診断": "", "改善案": "", "代替案": ""}
    current = None
    for line in improvement.splitlines():
        m = re.match(r'^###\s*(診断|改善案|代替案)', line)
        if m:
            current = m.group(1)
        elif current:
            sections[current] += line + "\n"
    # パースに失敗した場合は全文を改善案に入れる
    if not any(sections.values()):
        sections["改善案"] = improvement
    return sections


def add_question_block(doc: Document, question: str, improvement: str, idx: int):
    """1問分の色分けブロックを追加する"""
    _add_separator(doc)

    # 元の問題（グレー）
    _add_section_label(doc, f"元の問題　問{idx}", "D9D9D9")
    _add_content_lines(doc, question, "F5F5F5")

    sections = _parse_sections(improvement)

    # 診断（黄）
    _add_section_label(doc, "診断", "FFE699")
    _add_content_lines(doc, sections["診断"], "FFFDE7")

    # 改善案（青）
    _add_section_label(doc, "改善案", "9DC3E6")
    _add_content_lines(doc, sections["改善案"], "EBF4FF")

    # 代替案（緑）
    _add_section_label(doc, "代替案", "A9D18E")
    _add_content_lines(doc, sections["代替案"], "F0F7EE")


def build_word(results: list[dict], source_filename: str) -> Path:
    doc = Document()
    _set_page_margins(doc)

    # タイトル
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_line_spacing(p)
    run = p.add_run("テスト問題　改善レポート")
    _set_font(run, "游ゴシック", 14, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_line_spacing(p2)
    run2 = p2.add_run(
        f"元ファイル：{source_filename}　　{datetime.today().strftime('%Y年%m月%d日')}"
    )
    _set_font(run2, "游明朝", 10)

    for r in results:
        add_question_block(doc, r["question"], r["improvement"], r["idx"])

    timestamp = datetime.today().strftime("%Y%m%d_%H%M%S")
    out_path = OUTPUT_DIR / f"improve_{timestamp}.docx"
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# ================================================================
# メイン
# ================================================================

def main():
    sprint("=== テスト問題改善ツール ===")

    # 入力ファイルを探す
    IMPROVE_DIR.mkdir(parents=True, exist_ok=True)
    input_files = [p for p in IMPROVE_DIR.iterdir()
                   if p.suffix.lower() in {".docx", ".pdf", ".txt"}]
    if not input_files:
        sprint(f"[エラー] {IMPROVE_DIR} にファイルが見つかりません。")
        sprint("  Word(.docx) または PDF(.pdf) を置いてから実行してください。")
        sys.exit(1)

    target = input_files[0]
    sprint(f"対象ファイル: {unicodedata.normalize('NFC', target.name)}")

    # テキスト抽出
    raw_text = read_text_file(target)
    if not raw_text.strip():
        sprint("[エラー] ファイルからテキストを読み込めませんでした。")
        sprint("  スキャンされたPDFの場合は、テキストベースのPDFかWord形式に変換してください。")
        sys.exit(1)

    # 問題分割
    questions = parse_questions(raw_text)
    sprint(f"{len(questions)} 問を検出しました。")
    if not questions:
        sprint("[エラー] 問題を検出できませんでした。")
        sprint("  「問1」「問2」の形式で問題が記述されているか確認してください。")
        sys.exit(1)

    # 教科書分析（既存キャッシュから）
    textbook_analysis = ""
    try:
        cfg = load_config()
        ext_textbooks = cfg.get("sources", {}).get("external", {}).get("textbooks", [])
        for folder_str in ext_textbooks:
            folder = find_path(folder_str)
            if folder:
                analysis = get_or_create_analysis(folder)
                if analysis:
                    textbook_analysis += analysis + "\n\n"
        if textbook_analysis:
            sprint("教科書分析テキストを読み込みました。")
        else:
            sprint("[情報] 教科書分析テキストなしで実行します。")
    except Exception as e:
        sprint(f"[情報] 設定読み込みスキップ ({e})。教科書分析なしで実行します。")

    # 問題ごとに改善案を生成
    sprint(f"\n--- 改善案生成 ({len(questions)}問) ---")
    results = []
    for i, q in enumerate(questions, 1):
        improvement = improve_question(q, textbook_analysis, i)
        results.append({"idx": i, "question": q, "improvement": improvement})

    # Word 出力
    sprint("\nWordファイルを生成中...")
    out_path = build_word(results, unicodedata.normalize("NFC", target.name))
    sprint(f"\n完了！　→ {out_path}")


if __name__ == "__main__":
    main()
