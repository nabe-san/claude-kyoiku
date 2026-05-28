"""
テスト問題自動生成スクリプト（2フェーズ方式）

Phase 1: 教科書JPEG → Claude で構造化テキストに分析 → キャッシュ保存
Phase 2: 分析テキスト → Claude でテスト問題生成 → Word 出力

同じ教科書フォルダを使う場合、Phase 1 はキャッシュから読み込むため
2回目以降は画像送信なしで高速・低コストに動作する。
"""

import os, sys, json, base64, yaml, unicodedata
from pathlib import Path
from datetime import datetime
import anthropic
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv

load_dotenv()


def sprint(text: str):
    """NFC正規化してターミナルのエンコードエラーを回避する print"""
    normalized = unicodedata.normalize("NFC", str(text))
    print(normalized.encode(sys.stdout.encoding or "utf-8", errors="replace")
                    .decode(sys.stdout.encoding or "utf-8", errors="replace"))


BASE_DIR     = Path(__file__).parent
INPUT_DIR    = BASE_DIR / "input"
OUTPUT_DIR   = BASE_DIR / "output"
TEMPLATES_DIR = BASE_DIR / "templates"
CONFIG_PATH  = BASE_DIR / "exam_config.yaml"
CACHE_DIR    = BASE_DIR.parent / "教科書" / "分析キャッシュ"

MODEL      = "claude-sonnet-4-6"
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
TEXT_EXTS  = {".txt", ".pdf", ".docx"}


# ================================================================
# 共通ユーティリティ
# ================================================================

def load_config() -> dict:
    with open(CONFIG_PATH, encoding="utf-8") as f:
        return yaml.safe_load(f)


def find_path(path_str: str) -> Path | None:
    """
    NFC/NFD の違いでパスが見つからない場合に対応する。
    親フォルダをスキャンして NFC正規化した名前が一致するフォルダ/ファイルを返す。
    """
    path = Path(path_str)
    if path.exists():
        return path
    parent = path.parent
    target = unicodedata.normalize("NFC", path.name)
    if parent.exists():
        for child in parent.iterdir():
            if unicodedata.normalize("NFC", child.name) == target:
                return child
    return None


def read_text_file(path: Path) -> str:
    if path.suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    if path.suffix == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        except Exception as e:
            sprint(f"  [警告] PDF読み込み失敗: {path.name} ({e})")
            return ""
    if path.suffix == ".docx":
        try:
            doc = Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            sprint(f"  [警告] docx読み込み失敗: {path.name} ({e})")
            return ""
    return ""


# ================================================================
# Phase 1: 教科書画像 → 構造化テキスト分析
# ================================================================

ANALYSIS_PROMPT = """\
あなたは高校の歴史教科書の内容を分析するアシスタントです。
提示された教科書の画像を詳細に読み取り、以下の項目を構造化してまとめてください。
テスト問題を作成する際の根拠テキストとして使用するため、正確さを最優先にしてください。

## 出力項目

### 1. 主要概念・用語
重要語句とその意味・説明を箇条書きで列挙する。

### 2. 歴史的事実（年号・人物・出来事）
「西暦年：出来事（関係人物）」の形式で時系列に列挙する。

### 3. 因果関係
「原因 → 結果」の形式で整理する。

### 4. 歴史的意義・影響
各出来事が歴史の流れの中で持つ意味・影響を記述する。

### 5. グラフ・表・資料の内容
グラフや表がある場合は以下を記述する：
- 資料名（例：資料1「輸出品の割合」）
- 軸・凡例の説明
- 読み取れる数値・傾向（例：1885年の生糸の割合は約XX%）
- グラフから導ける歴史的結論

### 6. 地図・図版・風刺画
地図や図版がある場合は内容と歴史的文脈を記述する。

### 7. 一次史料・文字資料（枠囲み資料）【最重要】
教科書には四角で囲まれた番号タイトル付きの文字資料が掲載されている。
例：「6 キャラコ禁止法」「9 児童労働法（フランス、1841年）」「20 ガンジーがみた植民地支配下のインド」
これらを全て以下の形式で完全に転記する：

資料番号タイトル：【資料XX「〇〇」】
出典・注記：（括弧内の情報があれば）
本文：（枠内の文章を一字一句そのまま転記する）

※ 枠内の本文は省略せず、読み取れる限り完全に転記すること。
※ 読み取り不可の箇所は「（読み取り不可）」と明記する。

### 8. 教科書の「問い」
教科書本文中に掲載されている問いかけ（「問い」「考えよう」「まとめ」等のラベルがついた問題）を全て抽出する。
各問いの本文を完全に転記し、どのテーマ・節に対応するかを明記する。

## 注意事項
- 画像から読み取れない箇所は「（読み取り不可）」と明記する
- 数値が不鮮明・不確かな場合は推測せず「（数値不明）」と記す
- 教科書に書かれていない事実を補足・追加しない
- 一次史料（7番）は最も重要な出力項目である。枠囲み資料が1つでも見えたら必ず転記すること
"""


MAX_IMAGE_PX = 1568  # 多画像リクエスト時の安全な上限（API制限 2000px より小さく設定）

def load_images_from_folder(folder: Path) -> list[dict]:
    """フォルダ内の画像をリサイズして base64 エンコードして返す"""
    from PIL import Image
    import io
    blocks = []
    for p in sorted(folder.iterdir()):
        if p.suffix.lower() in IMAGE_EXTS:
            img = Image.open(p).convert("RGB")
            w, h = img.size
            if max(w, h) > MAX_IMAGE_PX:
                scale = MAX_IMAGE_PX / max(w, h)
                img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=75)
            data = base64.standard_b64encode(buf.getvalue()).decode()
            blocks.append({"type": "image",
                            "source": {"type": "base64", "media_type": "image/jpeg", "data": data}})
            sprint(f"    画像読み込み: {unicodedata.normalize('NFC', p.name)} ({img.size[0]}x{img.size[1]})")
    return blocks


def run_analysis(image_blocks: list[dict]) -> str:
    """Claude に画像を送り、構造化テキスト分析を返す"""
    client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
    content = list(image_blocks) + [{"type": "text", "text": ANALYSIS_PROMPT}]
    response = client.messages.create(
        model=MODEL,
        max_tokens=8192,
        messages=[{"role": "user", "content": content}],
    )
    return response.content[0].text


def get_or_create_analysis(folder: Path) -> str:
    """
    キャッシュがあれば読み込む。なければ画像を分析してキャッシュに保存する。
    キャッシュファイル: 教科書/分析キャッシュ/<フォルダ名>.json
    """
    folder_name = unicodedata.normalize("NFC", folder.name)
    cache_file = CACHE_DIR / f"{folder_name}.json"

    if cache_file.exists():
        sprint(f"  [キャッシュ使用] {folder_name}.json")
        data = json.loads(cache_file.read_text(encoding="utf-8"))
        return data["content"]

    # キャッシュなし → Phase 1 実行
    sprint(f"  [Phase 1] {folder_name} を分析中...")
    image_blocks = load_images_from_folder(folder)
    if not image_blocks:
        sprint(f"  [警告] 画像が見つかりません: {folder}")
        return ""

    analysis = run_analysis(image_blocks)

    CACHE_DIR.mkdir(parents=True, exist_ok=True)
    cache_data = {
        "folder": folder_name,
        "analyzed_at": datetime.now().isoformat(),
        "image_count": len(image_blocks),
        "content": analysis,
    }
    cache_file.write_text(json.dumps(cache_data, ensure_ascii=False, indent=2), encoding="utf-8")
    sprint(f"  [Phase 1 完了] キャッシュ保存 → {cache_file.name}")
    return analysis


# ================================================================
# 資料収集（テキストのみ・Phase 2 用）
# ================================================================

def collect_sources(cfg: dict) -> list[str]:
    """
    全資料をテキストとして収集する。
    教科書フォルダは get_or_create_analysis() 経由で分析テキストを取得。
    """
    text_chunks: list[str] = []

    def add_texts(folder: Path, label: str):
        for p in sorted(folder.iterdir()):
            if p.suffix.lower() in TEXT_EXTS:
                text = read_text_file(p)
                if text.strip():
                    name = unicodedata.normalize("NFC", p.name)
                    text_chunks.append(f"【{label}：{name}】\n{text.strip()}")
                    sprint(f"  テキスト読み込み: {name}")

    src = cfg.get("sources", {})

    # input/ 以下
    if src.get("use_past_exams"):
        add_texts(INPUT_DIR / "past_exams", "過去問")
    if src.get("use_worksheets"):
        add_texts(INPUT_DIR / "worksheets", "プリント")

    # 外部フォルダ
    ext = src.get("external", {})

    for folder_str in ext.get("past_exams", []):
        folder = find_path(folder_str)
        if folder:
            add_texts(folder, "過去問")
        else:
            sprint(f"  [警告] 過去問フォルダが見つかりません: {folder_str}")

    for folder_str in ext.get("textbooks", []):
        folder = find_path(folder_str)
        if folder:
            analysis = get_or_create_analysis(folder)
            if analysis:
                name = unicodedata.normalize("NFC", folder.name)
                text_chunks.append(f"【教科書分析：{name}】\n{analysis}")
        else:
            sprint(f"  [警告] 教科書フォルダが見つかりません: {folder_str}")

    for folder_str in ext.get("worksheets", []):
        folder = find_path(folder_str)
        if folder:
            add_texts(folder, "プリント")
        else:
            sprint(f"  [警告] プリントフォルダが見つかりません: {folder_str}")

    return text_chunks


# ================================================================
# Phase 2: テスト問題生成
# ================================================================

def build_prompt(cfg: dict, text_chunks: list[str]) -> str:
    q            = cfg["questions"]
    d            = cfg["difficulty"]
    exam         = cfg["exam"]
    instructions = cfg.get("instructions", "")
    points       = q.get("points_per_question", 1)
    total_score  = q["total"] * points
    source_block = "\n\n---\n\n".join(text_chunks) if text_chunks else "（参照資料なし）"

    n_primary   = q.get("primary_source", 6)
    n_toboi     = q.get("textbook_question", 8)
    n_timeline  = q.get("timeline_order", 2)
    n_other     = q.get("other", 4)

    return f"""\
あなたは高校の{exam['subject']}教員のアシスタントです。
以下の【教科書分析テキスト】を唯一の根拠として定期試験の問題を作成してください。
分析テキストに記載のない事実・年号・人名・因果関係を問題や選択肢に使わないでください。

## 試験設定
- 科目: {exam['subject']}　学年: {exam['grade']}　単元: {exam.get('unit') or '（指定なし）'}
- 問題数: 合計{q['total']}問　1問{points}点　満点{total_score}点
- 難易度: 標準{d['standard']}% / 難{d['hard']}%

## 問題の種類と必要数（厳守）

### ① 一次史料問題（{n_primary}問・全体の約30%）
教科書分析テキストの「### 7. 一次史料・文字資料」に記載された枠囲み資料を使う。
- 問題文中に資料本文（またはその一部）を直接引用する
- 形式例：
  「次の資料を読んで、問いに答えよ。
  資料6「キャラコ禁止法」（1721年、イギリス）
  第3条　1722年12月25日以降……（本文を引用）
  この法律が制定された背景として最も適切なものはどれか。」
- 選択肢は4択（ア〜エ）で、歴史的文脈の理解を要するものにする
- 資料が複数ある場合は異なる資料をそれぞれ1問ずつ使う

### ② 教科書「問い」問題（{n_toboi}問・全体の約40%）
教科書分析テキストの「### 8. 教科書の『問い』」に記載された問いを使う。
- 教科書の「問い」をそのまままたは4択化して出題する
- 問いが記述式の場合は4択に変換する（例：「〜について考えよ」→「〜の理由として最も適切なものはどれか」）
- 問いが見つからない場合は、教科書本文の重要な問いかけ・課題を4択問題として出題する

### ③ 年代整序問題（{n_timeline}問）
複数の出来事を古い順に並べる問題。
- 「次のア〜エの出来事を古い順に並べたものとして正しいものを選べ」形式

### ④ Web資料・概念問題等（{n_other}問）
- うち1問以上は、Web上で入手できる史料・グラフ・画像を使った問題にする
- 解説欄にURL・出典を必ず明記する
- 残りは因果関係・歴史的意義を問う概念問題

## 作成方針
{instructions}

## 教科書分析テキスト（この内容だけを根拠にすること）
{source_block}

---

## 出力形式（厳守）

全問題を出力する前に、以下の3点を自己確認すること：
1. 各問題の問題文・選択肢が自然な日本語として成立しているか
2. 各問題の正答の根拠が上記の教科書分析テキストから取れるか
3. ①〜④の種類と問題数が指定通りになっているか
→ 破綻・根拠不足・数の不足があれば出力前に書き直す

### 問題

**問{{番号}}**（{{①一次史料 or ②問い or ③年代整序 or ④概念/Web}}・{{標準 or 難}}）
{{問題文。一次史料問題は資料本文を引用する}}
ア．{{選択肢A}}
イ．{{選択肢B}}
ウ．{{選択肢C}}
エ．{{選択肢D}}

---

### 解答・解説

**問{{番号}}** 正答：{{正答記号}}
解説：{{教科書分析テキストのどの記述が根拠かを含む1〜2文。Web資料使用の場合はURL・出典を明記}}
"""


def call_claude(prompt: str) -> str:
    """Phase 2: テキストのみで問題生成（画像不要）"""
    client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
    sprint("\n[Phase 2] 問題を生成中...")
    response = client.messages.create(
        model=MODEL,
        max_tokens=8192,
        messages=[{"role": "user", "content": prompt}],
    )
    return response.content[0].text


# ================================================================
# Word 出力
# ================================================================

def parse_sections(raw: str) -> tuple[str, str]:
    if "### 解答・解説" in raw:
        parts = raw.split("### 解答・解説", 1)
        return parts[0], parts[1]
    return raw, ""


# ---- ページ設定 ----

def _set_page_margins(doc: Document):
    """A4縦・余白 上下20mm 左右25mm"""
    for section in doc.sections:
        section.page_width  = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)


# ---- 段落ヘルパー ----

def _set_line_spacing(para, multiple: float = 1.15):
    """行間を MULTIPLE で設定"""
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing      = multiple


def _set_font(run, name: str, size_pt: float, bold: bool = False,
              underline: bool = False):
    run.font.name  = name
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.underline = underline
    # 日本語フォントを東アジアフォントとして明示設定
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def _add_horizontal_rule(doc: Document):
    """問題間の水平罫線"""
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    pb.append(bottom)
    pPr.append(pb)
    _set_line_spacing(p)


# ---- スタイル分岐 ----

def add_styled_paragraph(doc: Document, line: str) -> None:
    """
    行の内容に応じてフォント・インデント・間隔を分岐して段落を追加する。

    分岐ルール:
      **問   → 游ゴシック・太字・11pt・前間隔12pt（問題番号行）
      ア．〜 → 游明朝・10.5pt・左インデント1cm（選択肢）
      資料   → 游ゴシック・太字・10.5pt（資料タイトル行）
      ---    → 水平罫線
      それ以外→ 游明朝・10.5pt
    """
    # 水平罫線マーカー
    if line.strip() in ("---", "―――"):
        _add_horizontal_rule(doc)
        return

    p = doc.add_paragraph()
    _set_line_spacing(p)

    clean = line.replace("**", "")

    # 問題番号行
    if line.startswith("**問") or line.startswith("問") and "（" in line:
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(clean)
        _set_font(run, "游ゴシック", 11, bold=True)

    # 選択肢行
    elif line.startswith(("ア．", "イ．", "ウ．", "エ．",
                           "ア.", "イ.", "ウ.", "エ.")):
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(clean)
        _set_font(run, "游明朝", 10.5)

    # 資料タイトル行
    elif "資料" in line:
        run = p.add_run(clean)
        _set_font(run, "游ゴシック", 10.5, bold=True)

    # 通常行
    else:
        run = p.add_run(clean)
        _set_font(run, "游明朝", 10.5)


def _add_title(doc: Document, text: str):
    """大見出し：游ゴシック・太字・14pt・中央揃え"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_line_spacing(p)
    run = p.add_run(text)
    _set_font(run, "游ゴシック", 14, bold=True)


def _add_meta(doc: Document, text: str):
    """副題・受験者情報行：游明朝・10.5pt・中央揃え"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_line_spacing(p)
    run = p.add_run(text)
    _set_font(run, "游明朝", 10.5)


def _add_answer_heading(doc: Document, text: str):
    """解答・解説ページの見出し：游ゴシック・太字・14pt"""
    p = doc.add_paragraph()
    _set_line_spacing(p)
    run = p.add_run(text)
    _set_font(run, "游ゴシック", 14, bold=True)


# ---- メイン出力関数 ----

def build_word(cfg: dict, raw_text: str) -> Path:
    doc = Document()
    _set_page_margins(doc)

    exam    = cfg["exam"]
    subject = exam["subject"]
    unit    = exam.get("unit") or ""
    grade   = exam["grade"]
    points  = cfg["questions"].get("points_per_question", 1)
    total   = cfg["questions"]["total"] * points

    # ヘッダー
    _add_title(doc, f"{subject}　定期試験")
    _add_meta(doc, f"{grade}　{datetime.today().strftime('%Y年%m月%d日')}")
    _add_meta(doc, f"氏名：＿＿＿＿＿＿＿＿＿＿　　番号：＿＿＿＿　　／{total}点")

    questions_text, answers_text = parse_sections(raw_text)

    # 問題本文
    for line in questions_text.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("###"):
            # セクション見出し（「問題」など）はスキップ or 小見出し
            continue
        add_styled_paragraph(doc, line)

    # 解答・解説シート（別ページ）
    if answers_text.strip() and cfg["output"].get("include_answer_sheet"):
        doc.add_page_break()
        _add_answer_heading(doc, "解答・解説")
        for line in answers_text.splitlines():
            line = line.strip()
            if not line:
                continue
            add_styled_paragraph(doc, line)

    # 保存
    base      = cfg["output"].get("filename", "テスト_{subject}_{unit}.docx")
    base      = base.format(subject=subject, unit=unit or "未設定").replace(".docx", "")
    timestamp = datetime.today().strftime("%Y%m%d_%H%M%S")
    out_path  = OUTPUT_DIR / f"{base}_{timestamp}.docx"
    OUTPUT_DIR.mkdir(exist_ok=True)
    doc.save(str(out_path))
    return out_path


# ================================================================
# メイン
# ================================================================

def main():
    sprint("=== テスト問題作成ツール（2フェーズ方式）===")
    cfg = load_config()
    unit = cfg["exam"].get("unit") or "単元未設定"
    sprint(f"設定: {cfg['exam']['subject']} / {cfg['exam']['grade']} / {unit}")

    sprint("\n--- Phase 1: 資料収集・教科書分析 ---")
    text_chunks = collect_sources(cfg)
    sprint(f"  → テキスト {len(text_chunks)} 件")

    if not text_chunks:
        sprint("\n[エラー] 資料が1件も読み込めませんでした。")
        sprint("  exam_config.yaml の sources 設定を確認してください。")
        sys.exit(1)

    sprint("\n--- Phase 2: 問題生成 ---")
    prompt = build_prompt(cfg, text_chunks)
    raw    = call_claude(prompt)

    sprint("\n--- Word ファイル生成 ---")
    out_path = build_word(cfg, raw)
    sprint(f"\n完了！　→ {out_path}")


if __name__ == "__main__":
    main()
