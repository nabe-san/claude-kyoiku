"""
教科書画像 → 一問一答カード 自動生成ツール（日本史・公共）

【フォルダ構造】
input/
  01_日露戦争/   ← 教科書画像を授業回ごとに入れる
    p001.jpg
    p002.jpg
  02_ビスマルク/
    ...
output/
  01_日露戦争/   ← input と同名フォルダに自動出力
    p001.docx
    p001_form_url.txt
  02_ビスマルク/
    ...

使い方:
  1. input/ の中に「授業回フォルダ」を作成（例: 03_明治維新）
  2. そのフォルダに教科書画像（JPEG/PNG）を入れる
  3. python generate_anki.py を実行
  4. output/ の同名フォルダに Word + Google Forms URL が出力される
"""

import os
import json
import re
import base64
import time
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

load_dotenv()
client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
MODEL = "claude-sonnet-4-6"

BASE_DIR   = Path(__file__).parent
INPUT_DIR  = BASE_DIR.parent / "教科書"
OUTPUT_DIR = BASE_DIR / "output"

IMAGE_EXTS     = {".jpg", ".jpeg", ".JPG", ".JPEG", ".png", ".PNG"}
WATCH_INTERVAL = 300  # 5分（秒）

CREDS_FILE    = BASE_DIR / "credentials.json"
TOKEN_FILE    = BASE_DIR / "token.json"
GOOGLE_SCOPES = ["https://www.googleapis.com/auth/forms.body"]

JPEG_MAGIC = b"\xff\xd8\xff"
PNG_MAGIC  = b"\x89PNG"


def _is_image(path: Path) -> bool:
    if path.suffix in IMAGE_EXTS:
        return True
    try:
        header = path.read_bytes()[:8]
        return header[:3] == JPEG_MAGIC or header[:4] == PNG_MAGIC
    except Exception:
        return False


def _media_type(path: Path) -> str:
    header = path.read_bytes()[:4]
    if header[:4] == PNG_MAGIC:
        return "image/png"
    return "image/jpeg"


# ─── 未処理画像の検出（授業回フォルダ対応）─────────────

def find_unprocessed() -> list[tuple[str, Path]]:
    """(session_name, image_path) のリストを返す。授業回フォルダのみ対象。"""
    results = []
    for session_dir in sorted(INPUT_DIR.iterdir()):
        if not session_dir.is_dir():
            continue
        session_name = session_dir.name
        out_session = OUTPUT_DIR / session_name
        processed = {p.stem for p in out_session.glob("*.docx")} if out_session.exists() else set()
        for img in sorted(session_dir.iterdir()):
            if img.is_file() and _is_image(img) and img.stem not in processed:
                results.append((session_name, img))
    return results


# ─── カード生成（画像 → Claude Vision）──────────────

CARD_SYSTEM = """\
あなたは高校日本史・公共の授業用フラッシュカードを作成するAIです。
提供された教科書の見開き画像から重要語句を空欄にした一問一答カードを生成します。

【ルール】
- 人名・地名・事件名・制度名・概念など重要語句を1つだけ空欄にする
- 空欄は「＿＿＿＿」と表記する
- 答えは空欄に入る語句のみ（1〜6語程度）
- **年号・西暦など「数字の年」を答えさせる問題は絶対に作らない**
- 20〜30枚作成する（少なすぎても多すぎても不可）
- 画像に写っている内容のみから出題する
- JSONのみ出力（前後の説明文は不要）

【出力形式】
[
  {"question": "○○は＿＿＿＿を制定した。", "answer": "大宝律令"},
  {"question": "＿＿＿＿は奴国の王に金印を授けた。", "answer": "後漢の光武帝"}
]
"""


def generate_cards(image_path: Path) -> list[dict]:
    media_type = _media_type(image_path)
    image_data = base64.standard_b64encode(image_path.read_bytes()).decode("utf-8")

    response = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        system=CARD_SYSTEM,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": media_type,
                        "data": image_data,
                    },
                },
                {
                    "type": "text",
                    "text": "この教科書の画像から一問一答カードを20〜30枚生成してください。",
                },
            ],
        }],
    )

    raw = response.content[0].text
    match = re.search(r"\[.*\]", raw, re.DOTALL)
    if not match:
        raise ValueError(f"JSONが見つかりませんでした:\n{raw}")
    return json.loads(match.group())


# ─── Word 保存 ────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def save_docx(cards: list[dict], output_path: Path, session_name: str, unit_name: str):
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"一問一答カード　{session_name}　{unit_name}")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x21, 0x37)

    doc.add_paragraph(f"全{len(cards)}問").runs[0].font.size = Pt(10)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Cm(11)
    table.columns[1].width = Cm(5)

    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["問題（表）", "答え（裏）"]):
        _set_cell_bg(cell, "0D2137")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)

    for i, card in enumerate(cards):
        row = table.add_row().cells
        bg = "F4F6F9" if i % 2 == 0 else "FFFFFF"
        _set_cell_bg(row[0], bg)
        _set_cell_bg(row[1], bg)

        q_run = row[0].paragraphs[0].add_run(f"{i+1}. {card['question']}")
        q_run.font.size = Pt(11)

        a_run = row[1].paragraphs[0].add_run(card["answer"])
        a_run.font.size = Pt(11)
        a_run.font.bold = True
        a_run.font.color.rgb = RGBColor(0xC0, 0x45, 0x08)

    doc.save(output_path)
    print(f"  → Word保存: output/{session_name}/{output_path.name}")


# ─── Google Forms 保存 ────────────────────────────

def _get_google_creds():
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import InstalledAppFlow
    from google.auth.transport.requests import Request

    creds = None
    if TOKEN_FILE.exists():
        creds = Credentials.from_authorized_user_file(str(TOKEN_FILE), GOOGLE_SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(str(CREDS_FILE), GOOGLE_SCOPES)
            creds = flow.run_local_server(port=0)
        TOKEN_FILE.write_text(creds.to_json(), encoding="utf-8")
    return creds


def create_google_form(cards: list[dict], session_name: str, unit_name: str) -> str:
    from googleapiclient.discovery import build

    creds = _get_google_creds()
    service = build("forms", "v1", credentials=creds)

    form = service.forms().create(body={
        "info": {"title": f"一問一答　{session_name}　{unit_name}"}
    }).execute()
    form_id = form["formId"]

    requests = [
        {
            "updateSettings": {
                "settings": {"quizSettings": {"isQuiz": True}},
                "updateMask": "quizSettings"
            }
        }
    ]
    for i, card in enumerate(cards):
        requests.append({
            "createItem": {
                "item": {
                    "title": f"{i+1}. {card['question']}",
                    "questionItem": {
                        "question": {
                            "required": False,
                            "grading": {
                                "pointValue": 1,
                                "correctAnswers": {
                                    "answers": [{"value": card["answer"]}]
                                }
                            },
                            "textQuestion": {"paragraph": False}
                        }
                    }
                },
                "location": {"index": i}
            }
        })

    service.forms().batchUpdate(
        formId=form_id,
        body={"requests": requests}
    ).execute()

    return f"https://docs.google.com/forms/d/{form_id}/viewform"


def save_form_url(url: str, output_path: Path, session_name: str, unit_name: str):
    output_path.write_text(
        f"一問一答　{session_name}　{unit_name}\nGoogle Forms URL:\n{url}\n",
        encoding="utf-8"
    )
    print(f"  → Forms URL保存: output/{session_name}/{output_path.name}")
    print(f"  → {url}")


# ─── 1画像の処理 ──────────────────────────────────

def process_image(session_name: str, image_path: Path):
    unit_name = image_path.stem
    out_dir = OUTPUT_DIR / session_name
    docx_path = out_dir / f"{unit_name}.docx"
    url_path  = out_dir / f"{unit_name}_form_url.txt"

    print(f"\n  [処理中] {session_name}/{image_path.name}")
    try:
        cards = generate_cards(image_path)
        print(f"  → {len(cards)}枚生成")
        save_docx(cards, docx_path, session_name, unit_name)
    except Exception as e:
        print(f"  [エラー] Word生成失敗 {image_path.name}: {e}")
        return

    if not CREDS_FILE.exists():
        print("  [スキップ] credentials.json がないため Google Forms 出力をスキップ")
        print("            セットアップ方法: README_FORMS_SETUP.txt を参照")
        return

    try:
        form_url = create_google_form(cards, session_name, unit_name)
        save_form_url(form_url, url_path, session_name, unit_name)
    except Exception as e:
        print(f"  [エラー] Google Forms生成失敗: {e}")


# ─── 監視ループ ───────────────────────────────────

def now() -> str:
    return datetime.now().strftime("%H:%M:%S")


def watch_loop():
    print("=" * 60)
    print("  教科書画像 → 一問一答カード 自動生成ツール")
    print()
    print("  【入力】教科書/<授業回フォルダ>/ に画像を入れる  ← プロジェクトルート直下")
    print("  【出力】output/<授業回フォルダ>/ に Word が出力される")
    print()
    print("  5分ごとに input/ を監視します。Ctrl+C で停止。")
    print("=" * 60)

    # 起動時に input/ のフォルダ一覧を表示
    session_dirs = sorted(d for d in INPUT_DIR.iterdir() if d.is_dir())
    if session_dirs:
        print(f"\n  検出した授業回フォルダ: {len(session_dirs)}件")
        for d in session_dirs:
            imgs = [f for f in d.iterdir() if f.is_file() and _is_image(f)]
            print(f"    {d.name}/  ({len(imgs)}枚の画像)")
    else:
        print(f"\n  [注意] 教科書/ に授業回フォルダがありません。")
        print(f"  例: 教科書/01_日露戦争/ フォルダを作成して画像を入れてください。")

    print()

    while True:
        items = find_unprocessed()
        if items:
            print(f"\n[{now()}] 未処理画像: {len(items)}件 → 処理開始")
            for session_name, img_path in items:
                process_image(session_name, img_path)
            print(f"\n[{now()}] 完了。次の確認まで5分待機...")
        else:
            print(f"[{now()}] 新しい画像なし。5分後に再確認します。")
        time.sleep(WATCH_INTERVAL)


# ─── メイン ─────────────────────────────────────────

def main():
    INPUT_DIR.mkdir(exist_ok=True)
    OUTPUT_DIR.mkdir(exist_ok=True)
    watch_loop()


if __name__ == "__main__":
    main()
