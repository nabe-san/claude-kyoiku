"""
授業音声（MP3） → 4択確認テスト 自動生成ツール
input/ の未処理 MP3 を5分ごとに監視し、Word を output/ に出力する
"""

import os
import json
import re
import base64
import time
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

load_dotenv()
client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
MODEL = "claude-sonnet-4-6"

BASE_DIR   = Path(__file__).parent
INPUT_DIR  = Path(r"G:\マイドライブ\生成AI【確認テスト】\input")
OUTPUT_DIR = Path(r"G:\マイドライブ\生成AI【確認テスト】\output")

AUDIO_EXTS     = {".mp3", ".MP3", ".m4a", ".M4A", ".wav", ".WAV"}
WATCH_INTERVAL = 300
MAX_MB         = 20


# ─── 未処理ファイルの検出 ─────────────────────────

def find_unprocessed() -> list[Path]:
    processed = {p.stem for p in OUTPUT_DIR.glob("*.docx")}
    files = [
        p for p in INPUT_DIR.iterdir()
        if p.suffix in AUDIO_EXTS and p.stem not in processed
    ]
    return sorted(files)


# ─── 4択問題生成 ─────────────────────────────────

QUIZ_SYSTEM = """\
あなたは高校日本史・公共の確認テストを作成するAIです。
提供された授業音声を聞いて、4択問題を生成します。

【ルール】
- 授業で説明された重要な概念・事実・因果関係から出題する
- 各問題に選択肢を4つ（正答1つ・誤答3つ）作る
- 誤答は紛らわしいが明確に間違いのある選択肢にする
- 15〜20問作成する
- JSONのみ出力（前後の説明文は不要）

【出力形式】
[
  {
    "question": "問題文",
    "choices": ["選択肢A", "選択肢B", "選択肢C", "選択肢D"],
    "answer": 0
  }
]
※ answer は正答の選択肢インデックス（0〜3）
"""


def generate_quiz(audio_path: Path) -> list[dict]:
    mb = audio_path.stat().st_size / 1024 / 1024
    if mb > MAX_MB:
        raise ValueError(f"ファイルサイズが{mb:.1f}MBで上限{MAX_MB}MBを超えています。")

    suffix = audio_path.suffix.lower()
    media_type = {"mp3": "audio/mpeg", "m4a": "audio/mp4", "wav": "audio/wav"}.get(
        suffix.lstrip("."), "audio/mpeg"
    )
    audio_data = base64.standard_b64encode(audio_path.read_bytes()).decode("utf-8")

    response = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        system=QUIZ_SYSTEM,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": media_type,
                        "data": audio_data,
                    },
                },
                {
                    "type": "text",
                    "text": "この授業音声から4択確認テストを15〜20問生成してください。",
                },
            ],
        }],
    )

    raw = response.content[0].text
    match = re.search(r"\[.*\]", raw, re.DOTALL)
    if not match:
        raise ValueError(f"JSONが見つかりませんでした:\n{raw}")
    return json.loads(match.group())


# ─── Word 保存 ────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


MARKS = ["①", "②", "③", "④"]


def save_docx(quiz: list[dict], output_path: Path, unit_name: str):
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # タイトル
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f"確認テスト　{unit_name}")
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x0D, 0x21, 0x37)
    doc.add_paragraph(f"全{len(quiz)}問　　氏名：＿＿＿＿＿＿＿＿").runs[0].font.size = Pt(10)

    # 問題
    for i, q in enumerate(quiz):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(f"問{i+1}．{q['question']}")
        run.font.size = Pt(12)
        run.font.bold = True

        for j, choice in enumerate(q["choices"]):
            cp = doc.add_paragraph()
            cp.paragraph_format.left_indent = Cm(1)
            cp.paragraph_format.space_before = Pt(2)
            cr = cp.add_run(f"{MARKS[j]}　{choice}")
            cr.font.size = Pt(11)

    # 答え合わせ表
    doc.add_page_break()
    ap = doc.add_paragraph()
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = ap.add_run("【解答】")
    ar.font.size = Pt(14)
    ar.font.bold = True
    ar.font.color.rgb = RGBColor(0x0D, 0x21, 0x37)

    cols = 5
    rows = -(-len(quiz) // cols)
    table = doc.add_table(rows=rows + 1, cols=cols * 2)
    table.style = "Table Grid"

    # ヘッダー
    for col in range(cols):
        for ci, label in enumerate(["問", "答"]):
            cell = table.rows[0].cells[col * 2 + ci]
            _set_cell_bg(cell, "0D2137")
            r = cell.paragraphs[0].add_run(label)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 解答行
    for idx, q in enumerate(quiz):
        row_i = idx // cols + 1
        col_i = idx % cols
        n_cell = table.rows[row_i].cells[col_i * 2]
        a_cell = table.rows[row_i].cells[col_i * 2 + 1]
        n_cell.paragraphs[0].add_run(str(idx + 1)).font.size = Pt(10)
        n_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar = a_cell.paragraphs[0].add_run(MARKS[q["answer"]])
        ar.font.size = Pt(10)
        ar.font.bold = True
        ar.font.color.rgb = RGBColor(0xC0, 0x45, 0x08)
        a_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    print(f"  → 保存: {output_path.name}")


# ─── 1ファイルの処理 ──────────────────────────────

def process_audio(audio_path: Path):
    unit_name = audio_path.stem
    output_path = OUTPUT_DIR / f"{unit_name}.docx"
    print(f"\n  [処理中] {audio_path.name}")
    try:
        quiz = generate_quiz(audio_path)
        print(f"  → {len(quiz)}問生成")
        save_docx(quiz, output_path, unit_name)
    except Exception as e:
        print(f"  [エラー] {audio_path.name}: {e}")


# ─── 監視ループ ───────────────────────────────────

def now() -> str:
    return datetime.now().strftime("%H:%M:%S")


def watch_loop():
    print("=" * 50)
    print("  授業音声 → 4択確認テスト 自動生成ツール")
    print("  5分ごとに input/ を監視します。Ctrl+C で停止。")
    print("=" * 50)
    while True:
        files = find_unprocessed()
        if files:
            print(f"\n[{now()}] 未処理: {len(files)}件 → 処理開始")
            for f in files:
                process_audio(f)
            print(f"\n[{now()}] 完了。次の確認まで5分待機...")
        else:
            print(f"[{now()}] 新しいファイルなし。5分後に再確認します。")
        time.sleep(WATCH_INTERVAL)


# ─── メイン ─────────────────────────────────────────

def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    watch_loop()


if __name__ == "__main__":
    main()
