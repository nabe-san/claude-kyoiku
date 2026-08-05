"""
プリント作成ツール — Word描画（Phase3）

structure.json（Phase2でClaudeが生成した構造化データ）を読み込み、
生徒作業用プリントと模範解答を「同じレイアウト関数群」から生成する。
両者の差分は解答欄の中身（空欄 or 記入済み）だけにすることで、
生成のたびにデザインがブレる問題を構造的になくしている。

デザイン（色・余白・罫線）を調整したいときはこのファイルだけを編集し、
Claude API を呼び直さずに再描画できる：
    python build_docx.py <教科書フォルダ名>

デザイン仕様は「歴史総合プリント作成：設計原則」（ケンゴさんがClaude.aiプロジェクトに
与えている指示書）に準拠する。色・フォント指定はそこからの移植。
"""

import sys
import json
import unicodedata
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).parent
OUTPUT_DIR = BASE_DIR / "output"

# ================================================================
# デザイン定数（色調整はここを編集する）
# ================================================================
# ネイビー・ブルー・ライトブルー・ゴールドは指示書の指定値をそのまま使用。
# 緑（解答例）だけは指示書に色指定がなく画像からの目視推定 — 要調整。

COLOR_NAVY = "1A3A5C"        # タイトル帯（大きな問い）・セクション帯
COLOR_BLUE = "2C5F8A"        # 見出し・ボーダー・補助資料ラベル
COLOR_LIGHT_BLUE = "D6E8F5"  # 用語ヘッダー・テーマ番号セル・補助資料背景
COLOR_GOLD = "7B5E00"        # ヒント文字・まとめ「使用語句」文字
COLOR_GOLD_BG = "FBF3DC"     # まとめ「使用語句」ボックス背景（推定）
COLOR_GREEN_HEADER = "1E6B3C"  # 解答例ラベル帯（推定・要調整）
COLOR_GREEN_BODY = "E3F2E7"    # 解答例本文背景（推定・要調整）
COLOR_RED = "C00000"         # 問N番号
COLOR_WHITE = "FFFFFF"
COLOR_BORDER_GRAY = "BFBFBF"
COLOR_TITLE_SHADE = "EDEDED"  # タイトル見出しの背景（Wordの見出し1相当のグレー）

FONT_MAIN = "BIZ UDPゴシック"       # 本文・見出し共通の基本フォント
FONT_BOX = "UD Digi Kyokasho NK-R"  # 補助資料・キーワード整理・解答例で使うフォント

LINE_SPACING_PT = 13  # 行間は倍率(1.15など)ではなく固定値(pt)を使う
BOX_FONT_SIZE = 9     # 解答例・キーワード整理・補助資料の文字サイズ
HINT_FONT_SIZE = 8    # ヒントの文字サイズ

CIRCLED_NUMBERS = ["①", "②", "③", "④", "⑤", "⑥"]

INVALID_FS_CHARS = '\\/:*?"<>|'


def sprint(text: str):
    """NFC正規化してターミナルのエンコードエラーを回避する print"""
    normalized = unicodedata.normalize("NFC", str(text))
    print(normalized.encode(sys.stdout.encoding or "utf-8", errors="replace")
                    .decode(sys.stdout.encoding or "utf-8", errors="replace"))


def _safe_filename(name: str) -> str:
    name = unicodedata.normalize("NFC", name)
    for ch in INVALID_FS_CHARS:
        name = name.replace(ch, "_")
    return name


def is_done(out_dir: Path) -> bool:
    """学生用・模範解答の両方が output/<フォルダ名>/ に揃っているか"""
    if not out_dir.exists():
        return False
    return any(out_dir.glob("*_学生用.docx")) and any(out_dir.glob("*_模範解答.docx"))


# ================================================================
# 低レベルヘルパー
# ================================================================

def _set_page_margins(doc: Document):
    """A4縦・余白を狭めにしてページ幅を目一杯使う"""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)


def _set_font(run, name: str, size_pt: float, bold: bool = False, color: str = None):
    """
    東アジアフォントのみ指定する（指示書のXML仕様に準拠）。
    w:ascii/w:hAnsi/w:cs は設定せず、w:eastAsia + w:hint="eastAsia" のみを設定する。
    """
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

    rPr = run._r.get_or_add_rPr()
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:hint"), "eastAsia")
    rPr.insert(0, rFonts)


def _set_line_spacing(para, pt: float = LINE_SPACING_PT):
    """行間は倍率ではなく固定値(pt)で設定する"""
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(pt)


def _spacer(doc: Document, pt: float = 6):
    """要素間の余白を最小限の高さで確保する（既定の空段落だと間延びするため）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(pt)


def _set_cell_bg(cell, hex_color: str):
    """セル背景色を設定する（授業自動生成/generate.py の実装を移植）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _no_border(cell):
    """セルの罫線を消す（帯の外枠を消して色面だけにする）"""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tcPr.append(borders)


def _cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """セル内余白を設定する（単位: twips = 1/20pt）"""
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def _set_row_cantsplit(row):
    """行がページをまたいで分割されないようにする"""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def _set_row_height(row, lines: int, twips_per_line: int = 420):
    """行の高さを明示指定する（記入欄の行数確保用。単位: twips）"""
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(max(lines, 1) * twips_per_line))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def _printable_width(doc: Document) -> int:
    section = doc.sections[0]
    return section.page_width - section.left_margin - section.right_margin


def _full_width_box(doc: Document):
    """本文幅いっぱいの1×1 Table Grid を作り、(table, cell) を返す。cantSplit済み"""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = _printable_width(doc)
    cell = table.cell(0, 0)
    cell.width = _printable_width(doc)
    _set_row_cantsplit(table.rows[0])
    return table, cell


# ================================================================
# レイアウト部品
# ================================================================

def add_plain_title(doc: Document, title: str):
    """単元タイトル：塗りつぶしなしの見出し（左に太い縦罫線＋下線）"""
    p = doc.add_paragraph()
    _set_line_spacing(p)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.2)

    pPr = p._p.get_or_add_pPr()

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), COLOR_TITLE_SHADE)
    pPr.append(shd)

    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), COLOR_NAVY)
    pBdr.append(left)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)

    _set_font(p.add_run(f"　{title}"), FONT_MAIN, 16, bold=True)
    _spacer(doc)


def add_big_question_banner(doc: Document, big_question: str):
    """大きな問い：ネイビー帯＋白太字"""
    table, cell = _full_width_box(doc)
    _set_cell_bg(cell, COLOR_NAVY)
    _no_border(cell)
    _cell_margins(cell, top=40, bottom=40, left=150, right=150)
    p = cell.paragraphs[0]
    _set_line_spacing(p)
    _set_font(p.add_run(big_question), FONT_MAIN, 11.5, bold=True, color=COLOR_WHITE)
    _spacer(doc)


def add_overview_table(doc: Document, questions: list[str]):
    """テーマへのアプローチ：丸数字（ライトブルー）＋テーマ文"""
    intro = doc.add_paragraph()
    _set_line_spacing(intro)
    _set_font(
        intro.add_run(f"【テーマへのアプローチ】　このプリントでは次の{len(questions)}つの問いを軸に考えていこう。"),
        FONT_MAIN, 10.5, bold=True,
    )

    table = doc.add_table(rows=len(questions), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    num_col_width = Cm(1.4)
    text_col_width = _printable_width(doc) - num_col_width
    table.columns[0].width = num_col_width
    table.columns[1].width = text_col_width

    for i, question_text in enumerate(questions):
        row = table.rows[i]
        _set_row_cantsplit(row)
        num_cell, text_cell = row.cells
        num_cell.width = num_col_width
        text_cell.width = text_col_width

        _set_cell_bg(num_cell, COLOR_LIGHT_BLUE)
        num_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_num = num_cell.paragraphs[0]
        p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mark = CIRCLED_NUMBERS[i] if i < len(CIRCLED_NUMBERS) else str(i + 1)
        _set_font(p_num.add_run(mark), FONT_MAIN, 12, bold=True)

        p_text = text_cell.paragraphs[0]
        _set_line_spacing(p_text)
        _set_font(p_text.add_run(question_text), FONT_MAIN, 10.5)

    _spacer(doc)


def add_section_banner(doc: Document, roman: str, heading: str, subtitle: str = ""):
    """セクション帯：ネイビー地に白太字。roman heading ─── subtitle"""
    table, cell = _full_width_box(doc)
    _set_cell_bg(cell, COLOR_NAVY)
    _no_border(cell)
    _cell_margins(cell, top=40, bottom=40, left=150, right=150)

    p = cell.paragraphs[0]
    _set_line_spacing(p)
    _set_font(p.add_run(f"{roman}　{heading}"), FONT_MAIN, 12.5, bold=True, color=COLOR_WHITE)
    if subtitle:
        _set_font(p.add_run(f"　───────　{subtitle}"), FONT_MAIN, 10.5, bold=True, color=COLOR_WHITE)

    _spacer(doc)


def add_glossary(doc: Document, label: str | None, terms: list[dict]):
    """キーワード整理：用語（ライトブルー太字）＋説明の2列表"""
    if label:
        p = doc.add_paragraph()
        _set_line_spacing(p)
        _set_font(p.add_run(label), FONT_BOX, BOX_FONT_SIZE, bold=True, color=COLOR_BLUE)

    term_col_width = Cm(2.8)
    text_col_width = _printable_width(doc) - term_col_width
    table = doc.add_table(rows=len(terms), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = term_col_width
    table.columns[1].width = text_col_width

    for i, entry in enumerate(terms):
        row = table.rows[i]
        _set_row_cantsplit(row)
        term_cell, def_cell = row.cells
        term_cell.width = term_col_width
        def_cell.width = text_col_width
        _cell_margins(term_cell, top=30, bottom=30, left=80, right=80)
        _cell_margins(def_cell, top=30, bottom=30, left=100, right=100)

        _set_cell_bg(term_cell, COLOR_LIGHT_BLUE)
        term_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_term = term_cell.paragraphs[0]
        p_term.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p_term.add_run(entry["term"]), FONT_BOX, BOX_FONT_SIZE, bold=True)

        p_def = def_cell.paragraphs[0]
        _set_line_spacing(p_def)
        _set_font(p_def.add_run(entry["definition"]), FONT_BOX, BOX_FONT_SIZE)

    _spacer(doc)


def add_supplementary_box(doc: Document, label: str, body: str):
    """補助資料：ライトブルー背景。ラベルはブルー太字"""
    table, cell = _full_width_box(doc)
    _set_cell_bg(cell, COLOR_LIGHT_BLUE)
    _cell_margins(cell, top=50, bottom=50, left=150, right=150)

    p1 = cell.paragraphs[0]
    _set_line_spacing(p1)
    _set_font(p1.add_run(label), FONT_BOX, BOX_FONT_SIZE, bold=True, color=COLOR_BLUE)

    p2 = cell.add_paragraph()
    _set_line_spacing(p2)
    _set_font(p2.add_run(body), FONT_BOX, BOX_FONT_SIZE)

    _spacer(doc)


def add_question(doc: Document, number: int, text: str, reference: str = None):
    """問N（赤太字）＋設問文。reference は教科書の問いに資料番号が無い場合のみ補足"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    _set_line_spacing(p)
    _set_font(p.add_run(f"問{number}　"), FONT_MAIN, 11, bold=True, color=COLOR_RED)
    full_text = text if not reference else f"{text}　（→{reference}を参照）"
    _set_font(p.add_run(full_text), FONT_MAIN, 11)


def add_hint(doc: Document, hint_text: str):
    """ヒント：ゴールド文字。難しい問いにのみ付ける"""
    p = doc.add_paragraph()
    _set_line_spacing(p)
    _set_font(p.add_run(f"💡 ヒント：{hint_text}"), FONT_MAIN, HINT_FONT_SIZE, color=COLOR_GOLD)


def add_answer_box(doc: Document, filled: bool, answer: str = None, point: str = None, blank_lines: int = 3):
    """
    解答欄。
    filled=False（学生用）: 罫線のみの1行テーブル。trHeightで行数分の記入スペースを確保。
    filled=True（模範解答）: 「解答例」ラベル帯（緑濃色）＋本文（緑薄色、▶ポイント付き）。
    """
    if filled:
        header_table, header_cell = _full_width_box(doc)
        _set_cell_bg(header_cell, COLOR_GREEN_HEADER)
        _no_border(header_cell)
        _cell_margins(header_cell, top=30, bottom=30, left=150, right=150)
        p_label = header_cell.paragraphs[0]
        _set_font(p_label.add_run("解答例"), FONT_BOX, BOX_FONT_SIZE, bold=True, color=COLOR_WHITE)

        # 本文は罫線を残す（キーワード整理・補助資料の箱と縦のラインを揃えるため）
        body_table, body_cell = _full_width_box(doc)
        _set_cell_bg(body_cell, COLOR_GREEN_BODY)
        _cell_margins(body_cell, top=50, bottom=50, left=150, right=150)
        p1 = body_cell.paragraphs[0]
        _set_line_spacing(p1)
        _set_font(p1.add_run(answer or ""), FONT_BOX, BOX_FONT_SIZE)
        if point:
            p2 = body_cell.add_paragraph()
            _set_line_spacing(p2)
            _set_font(p2.add_run(f"▶ ポイント：{point}"), FONT_BOX, BOX_FONT_SIZE, bold=True)
    else:
        table, cell = _full_width_box(doc)
        _cell_margins(cell, top=100, bottom=100, left=150, right=150)
        _set_row_height(table.rows[0], blank_lines)

    _spacer(doc)


def add_comparison_question(doc: Document, number: int, text: str, columns: list[str], hint: str = None):
    """比較設問：問N＋設問文＋2(以上)列の見出し行（ライトブルー）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    _set_line_spacing(p)
    _set_font(p.add_run(f"問{number}　"), FONT_MAIN, 11, bold=True, color=COLOR_RED)
    _set_font(p.add_run(text), FONT_MAIN, 11)
    if hint:
        add_hint(doc, hint)

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.autofit = False
    col_width = _printable_width(doc) // len(columns)
    for col in table.columns:
        col.width = col_width
    _set_row_cantsplit(table.rows[0])

    for cell, label in zip(table.rows[0].cells, columns):
        cell.width = col_width
        _set_cell_bg(cell, COLOR_LIGHT_BLUE)
        p_label = cell.paragraphs[0]
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p_label.add_run(label), FONT_MAIN, 10.5, bold=True)


def add_comparison_answer_box(doc: Document, columns: list[str], answers: list, filled: bool, blank_lines: int = 4):
    """比較設問の解答欄：列ごとに空欄 or 模範解答（緑背景）"""
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.autofit = False
    col_width = _printable_width(doc) // len(columns)
    for col in table.columns:
        col.width = col_width
    row = table.rows[0]
    _set_row_cantsplit(row)
    if not filled:
        _set_row_height(row, blank_lines)

    for i, cell in enumerate(row.cells):
        cell.width = col_width
        _cell_margins(cell, top=50, bottom=50, left=100, right=100)
        if filled:
            _set_cell_bg(cell, COLOR_GREEN_BODY)
            answer_text = answers[i] if i < len(answers) and answers[i] else ""
            p = cell.paragraphs[0]
            _set_line_spacing(p)
            _set_font(p.add_run(answer_text), FONT_BOX, BOX_FONT_SIZE)

    _spacer(doc)


def add_summary(doc: Document, central_question: str, vocabulary: list[str], answer: str, point: str, filled: bool):
    """まとめ：中心の問い（プレーン太字）＋使用語句（ゴールド枠）＋解答例"""
    p = doc.add_paragraph()
    _set_line_spacing(p)
    _set_font(p.add_run(f"【中心の問い】　{central_question}"), FONT_MAIN, 11, bold=True)
    _spacer(doc)

    table, cell = _full_width_box(doc)
    _set_cell_bg(cell, COLOR_GOLD_BG)
    _cell_margins(cell, top=50, bottom=50, left=150, right=150)
    p_vocab = cell.paragraphs[0]
    _set_line_spacing(p_vocab)
    _set_font(p_vocab.add_run("【使用語句】　" + "／".join(vocabulary)), FONT_MAIN, 10.5, bold=True, color=COLOR_GOLD)
    _spacer(doc)

    add_answer_box(doc, filled=filled, answer=answer, point=point)


def add_footer_note(doc: Document, subject: str, title: str, page_range: str):
    p = doc.add_paragraph()
    _set_line_spacing(p)
    _set_font(p.add_run(f"{subject}　授業プリント　{title}　（教科書 {page_range} 使用）"), FONT_MAIN, 9)


# ================================================================
# 文書全体の組み立て
# ================================================================

def render_document(structure: dict, filled: bool) -> Document:
    doc = Document()
    _set_page_margins(doc)

    add_plain_title(doc, structure["title"])
    add_big_question_banner(doc, structure["big_question"])
    add_overview_table(doc, structure["overview_questions"])

    for section in structure["sections"]:
        add_section_banner(doc, section["roman"], section["heading"], section.get("subtitle", ""))
        for item in section["items"]:
            t = item["type"]
            if t == "glossary":
                add_glossary(doc, item.get("label"), item["terms"])
            elif t == "supplementary":
                add_supplementary_box(doc, item["label"], item["body"])
            elif t == "question":
                add_question(doc, item["number"], item["text"], item.get("reference"))
                if item.get("hint"):
                    add_hint(doc, item["hint"])
                add_answer_box(doc, filled=filled, answer=item.get("answer"), point=item.get("point"))
            elif t == "comparison_question":
                add_comparison_question(doc, item["number"], item["text"], item["columns"], item.get("hint"))
                add_comparison_answer_box(doc, item["columns"], item.get("answers", []), filled=filled)
            elif t == "summary":
                add_summary(
                    doc,
                    item["central_question"],
                    item["vocabulary"],
                    item.get("answer"),
                    item.get("point"),
                    filled=filled,
                )

    if structure.get("page_range"):
        add_footer_note(doc, structure.get("subject", "歴史総合"), structure["title"], structure["page_range"])

    return doc


def build_student_docx(structure: dict, out_path: Path):
    doc = render_document(structure, filled=False)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    sprint(f"  → {out_path.name} 保存")


def build_answer_docx(structure: dict, out_path: Path):
    doc = render_document(structure, filled=True)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    sprint(f"  → {out_path.name} 保存")


def build_both(structure: dict, out_dir: Path):
    """structure.json から学生用・模範解答の2ファイルを out_dir に生成する"""
    title = _safe_filename(structure["title"])
    build_student_docx(structure, out_dir / f"{title}_学生用.docx")
    build_answer_docx(structure, out_dir / f"{title}_模範解答.docx")


# ================================================================
# CLI（structure.json からデザインだけ再描画したいときに使う）
# ================================================================

def main():
    if len(sys.argv) < 2:
        sprint("使い方: python build_docx.py <教科書フォルダ名>")
        sprint("  例:   python build_docx.py 第6回_変容する東アジア")
        sprint("  ※ 先に generate.py を実行して output/<フォルダ名>/structure.json を作成しておくこと")
        sys.exit(1)

    folder_name = unicodedata.normalize("NFC", sys.argv[1])
    out_dir = OUTPUT_DIR / folder_name
    structure_path = out_dir / "structure.json"
    if not structure_path.exists():
        sprint(f"[エラー] {structure_path} が見つかりません。先に generate.py を実行してください。")
        sys.exit(1)

    structure = json.loads(structure_path.read_text(encoding="utf-8"))
    sprint(f"[再描画] {structure_path} → 学生用・模範解答 docx")
    build_both(structure, out_dir)


if __name__ == "__main__":
    main()
