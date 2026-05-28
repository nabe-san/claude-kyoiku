#!/usr/bin/env python
"""
授業自動生成ツール

教科書の授業回ページ（画像 4 枚程度）を input/<授業回名>/ に入れて
python generate.py を実行すると、以下を output/<授業回名>/ に自動生成する。

  一問一答カード.docx
  テスト問題.docx

処理フロー:
  Phase 1  : Claude Vision で全画像を一括分析 → analysis.txt（2 ツール共有・キャッシュ）
  Phase 2a : 分析テキスト → 一問一答カード（Word）
  Phase 2b : 分析テキスト → テスト問題（Word）
"""

import os, sys, json, base64, re, unicodedata, io
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()

import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    print("[警告] Pillow 未インストール。pip install Pillow を実行してください。画像リサイズなしで処理します。")

client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
MODEL = "claude-sonnet-4-6"

BASE_DIR   = Path(__file__).parent
INPUT_DIR  = BASE_DIR / "input"
OUTPUT_DIR = BASE_DIR / "output"
BOOKS_DIR  = BASE_DIR.parent / "歴史シミュレーション" / "books"

IMAGE_EXTS   = {".jpg", ".jpeg", ".png", ".webp"}
MAX_IMAGE_PX = 1568
JPEG_MAGIC   = b"\xff\xd8\xff"
PNG_MAGIC    = b"\x89PNG"


# ================================================================
# ユーティリティ
# ================================================================

def sprint(text: str):
    normalized = unicodedata.normalize("NFC", str(text))
    print(normalized.encode(sys.stdout.encoding or "utf-8", errors="replace")
                    .decode(sys.stdout.encoding or "utf-8", errors="replace"))


def _is_image(path: Path) -> bool:
    if path.suffix.lower() in IMAGE_EXTS:
        return True
    try:
        header = path.read_bytes()[:8]
        return header[:3] == JPEG_MAGIC or header[:4] == PNG_MAGIC
    except Exception:
        return False


def load_images(folder: Path) -> list[dict]:
    """画像フォルダ → Claude Vision 用 content block リスト"""
    blocks = []
    for p in sorted(folder.iterdir()):
        if not (p.is_file() and _is_image(p)):
            continue
        if PIL_AVAILABLE:
            img = Image.open(p).convert("RGB")
            w, h = img.size
            if max(w, h) > MAX_IMAGE_PX:
                scale = MAX_IMAGE_PX / max(w, h)
                img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=75)
            data = base64.standard_b64encode(buf.getvalue()).decode()
            media_type = "image/jpeg"
        else:
            raw = p.read_bytes()
            data = base64.standard_b64encode(raw).decode()
            media_type = "image/png" if raw[:4] == PNG_MAGIC else "image/jpeg"
        blocks.append({
            "type": "image",
            "source": {"type": "base64", "media_type": media_type, "data": data}
        })
        sprint(f"    画像: {unicodedata.normalize('NFC', p.name)}")
    return blocks


# ================================================================
# Phase 1 : 教科書分析（3 ツール共有）
# ================================================================

ANALYSIS_PROMPT = """\
以下の教科書画像を詳細に分析し、構造化テキストにまとめてください。
一問一答カード・テスト問題の共通素材として使います。

### 1. 単元情報
- 授業回タイトル・単元名
- 時代・地域・主な人物

### 2. 主要概念・用語
重要語句とその意味・説明を箇条書きで列挙。

### 3. 歴史的事実（年号・人物・出来事）
「西暦年：出来事（関係人物）」の形式で時系列に列挙。

### 4. 因果関係
「原因 → 結果」の形式で整理。

### 5. 歴史的意義・影響
各出来事が歴史の流れの中で持つ意味・影響。

### 6. グラフ・表・資料の内容
資料名・軸・凡例・読み取れる数値・傾向・歴史的結論。

### 7. 一次史料・文字資料（枠囲み資料）【最重要】
四角で囲まれた番号タイトル付き文字資料を全て転記。
形式: 【資料XX「〇〇」】出典 / 本文（一字一句そのまま転記）
読み取り不可の箇所は「（読み取り不可）」と明記。

### 8. 教科書の「問い」
教科書本文中の問いかけ（「問い」「考えよう」等のラベルがある問題）を全て転記。

注意: 教科書に書かれていない事実を補足しない。数値不明な場合は「（数値不明）」と記す。
"""


def analyze_images(image_blocks: list[dict]) -> str:
    content = list(image_blocks) + [{"type": "text", "text": ANALYSIS_PROMPT}]
    response = client.messages.create(
        model=MODEL,
        max_tokens=8192,
        messages=[{"role": "user", "content": content}]
    )
    return response.content[0].text


# ================================================================
# Phase 2a : 一問一答カード → Word
# ================================================================

ANKI_SYSTEM = """\
あなたは高校歴史の授業用フラッシュカードを作成するAIです。
教科書分析テキストから重要語句を空欄にした一問一答カードを生成します。

ルール:
- 人名・地名・事件名・制度名・概念など重要語句を1つだけ空欄にする
- 空欄は「＿＿＿＿」と表記する
- 答えは空欄に入る語句のみ（1〜6語程度）
- 年号・西暦など「数字の年」を答えさせる問題は絶対に作らない
- 20〜30枚作成する
- JSONのみ出力（前後の説明文不要）

出力形式:
[
  {"question": "○○は＿＿＿＿を制定した。", "answer": "大宝律令"},
  {"question": "＿＿＿＿は奴国の王に金印を授けた。", "answer": "後漢の光武帝"}
]
"""


def generate_anki_cards(analysis: str) -> list[dict]:
    response = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        system=ANKI_SYSTEM,
        messages=[{
            "role": "user",
            "content": f"以下の教科書分析テキストから一問一答カードを20〜30枚生成してください。\n\n{analysis}"
        }]
    )
    raw = response.content[0].text
    match = re.search(r"\[.*\]", raw, re.DOTALL)
    if not match:
        raise ValueError(f"JSONが見つかりません:\n{raw[:200]}")
    return json.loads(match.group())


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def save_anki_docx(cards: list[dict], output_path: Path, session_name: str):
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"一問一答カード　{session_name}")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x21, 0x37)

    doc.add_paragraph(f"全{len(cards)}問").runs[0].font.size = Pt(10)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Cm(11)
    table.columns[1].width = Cm(5)

    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["問題（表）", "答え（裏）"]):
        _set_cell_bg(cell, "0D2137")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(text)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(11)

    for i, card in enumerate(cards):
        row = table.add_row().cells
        bg = "F4F6F9" if i % 2 == 0 else "FFFFFF"
        _set_cell_bg(row[0], bg)
        _set_cell_bg(row[1], bg)
        q_run = row[0].paragraphs[0].add_run(f"{i+1}. {card['question']}")
        q_run.font.size = Pt(11)
        a_run = row[1].paragraphs[0].add_run(card["answer"])
        a_run.font.size = Pt(11)
        a_run.font.bold = True
        a_run.font.color.rgb = RGBColor(0xC0, 0x45, 0x08)

    doc.save(output_path)
    sprint(f"  → 一問一答カード.docx 保存")


# ================================================================
# Phase 2b : テスト問題 → Word
# ================================================================

TEST_PROMPT_TEMPLATE = """\
以下の教科書分析テキストをもとに、高校歴史の定期試験問題を作成してください。
問題文は高校生が読む試験問題として完結した自然な日本語にする。
「教科書分析テキスト」「分析テキストによれば」などの表現は絶対に使わない。

## 試験設定
- 問題数: 20問（4択・年代整序のみ。論述・語句記述は作らない）
- 難易度: 標準 20% / 難しい 80%（易しい問題は作らない）

## 問題の種類と必要数
①一次史料問題（6問）: 分析テキスト「7. 一次史料」の枠囲み資料を本文引用して出題
②教科書「問い」問題（8問）: 分析テキスト「8. 教科書の問い」を4択化して出題
③年代整序問題（2問）: 複数の出来事を古い順に並べる問題
④概念・因果問題（4問）: 歴史的意義・因果関係・背景を問う。うち1問以上はWeb上の史料・グラフを使う（解説欄にURL・出典を明記）

## 作成方針
- 選択肢は4択（ア〜エ）。長さは30字程度
- 概念的知識（因果・意義・背景）を問う問題を優先
- 問題間で問題文・選択肢・正答が被らないようにする
- 全問生成後にセルフチェック：①日本語の自然さ ②根拠の確認 ③不自然な表現の除去

## 教科書分析テキスト
{analysis}

---

## 出力形式（厳守）

### 問題

**問{{番号}}**（{{①〜④の種別}}・{{標準 or 難}}）
{{問題文。一次史料問題は資料本文を引用する}}
ア．{{選択肢A}}
イ．{{選択肢B}}
ウ．{{選択肢C}}
エ．{{選択肢D}}

---

### 解答・解説

**問{{番号}}** 正答：{{正答記号}}
解説：{{根拠となる教科書の記述を含む1〜2文。Web資料使用の場合はURL・出典を明記}}
"""


def generate_test_questions(analysis: str) -> str:
    prompt = TEST_PROMPT_TEMPLATE.format(analysis=analysis)
    response = client.messages.create(
        model=MODEL,
        max_tokens=8192,
        messages=[{"role": "user", "content": prompt}]
    )
    return response.content[0].text


def _set_line_spacing(para, multiple: float = 1.15):
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = multiple


def _set_font(run, name: str, size_pt: float, bold: bool = False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def _add_hr(doc: Document):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    pb.append(bottom)
    pPr.append(pb)
    _set_line_spacing(para)


def save_test_docx(raw_text: str, output_path: Path, session_name: str):
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for section in doc.sections:
        section.page_width    = Cm(21.0)
        section.page_height   = Cm(29.7)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    for text, size, bold, align in [
        ("歴史　定期試験", 14, True, WD_ALIGN_PARAGRAPH.CENTER),
        (f"{session_name}　{datetime.today().strftime('%Y年%m月%d日')}", 10.5, False, WD_ALIGN_PARAGRAPH.CENTER),
        ("氏名：＿＿＿＿＿＿＿＿　番号：＿＿＿　／40点", 10.5, False, WD_ALIGN_PARAGRAPH.CENTER),
    ]:
        para = doc.add_paragraph()
        para.alignment = align
        _set_line_spacing(para)
        run = para.add_run(text)
        font_name = "游ゴシック" if bold else "游明朝"
        _set_font(run, font_name, size, bold)

    questions_text, answers_text = (
        raw_text.split("### 解答・解説", 1)
        if "### 解答・解説" in raw_text
        else (raw_text, "")
    )

    for line in questions_text.splitlines():
        line = line.strip()
        if not line or line.startswith("###"):
            continue
        if line.strip() in ("---", "―――"):
            _add_hr(doc)
            continue
        para = doc.add_paragraph()
        _set_line_spacing(para)
        clean = line.replace("**", "")
        if line.startswith("**問") or (line.startswith("問") and "（" in line):
            para.paragraph_format.space_before = Pt(12)
            run = para.add_run(clean)
            _set_font(run, "游ゴシック", 11, bold=True)
        elif line[:2] in ("ア．", "イ．", "ウ．", "エ．", "ア.", "イ.", "ウ.", "エ."):
            para.paragraph_format.left_indent = Cm(1.0)
            run = para.add_run(clean)
            _set_font(run, "游明朝", 10.5)
        elif "資料" in line:
            run = para.add_run(clean)
            _set_font(run, "游ゴシック", 10.5, bold=True)
        else:
            run = para.add_run(clean)
            _set_font(run, "游明朝", 10.5)

    if answers_text.strip():
        doc.add_page_break()
        para = doc.add_paragraph()
        _set_line_spacing(para)
        run = para.add_run("解答・解説")
        _set_font(run, "游ゴシック", 14, bold=True)
        for line in answers_text.splitlines():
            line = line.strip()
            if not line:
                continue
            para = doc.add_paragraph()
            _set_line_spacing(para)
            run = para.add_run(line.replace("**", ""))
            _set_font(run, "游明朝", 10.5)

    doc.save(output_path)
    sprint(f"  → テスト問題.docx 保存")


# ================================================================
# Phase 2c : 文字資料（行間補完）→ Word
# ================================================================

def load_book_excerpts(session_name: str, max_chars: int = 4000) -> str:
    """セッション名に関連する books/ の抜粋を返す。"""
    if not BOOKS_DIR.exists():
        return ""
    clean = re.sub(r"[\d第回_\s　]+", " ", session_name)
    keywords = [w for w in clean.split() if len(w) >= 2]
    if not keywords:
        return ""

    excerpts = []
    total = 0
    for book_path in sorted(BOOKS_DIR.glob("*.txt")):
        try:
            lines = book_path.read_text(encoding="utf-8", errors="replace").split("\n")
        except Exception:
            continue
        seen = set()
        for i, line in enumerate(lines):
            if not any(kw in line for kw in keywords):
                continue
            start, end = max(0, i - 2), min(len(lines), i + 9)
            chunk = "\n".join(lines[start:end]).strip()
            if len(chunk) < 30 or chunk in seen:
                continue
            seen.add(chunk)
            stem = book_path.stem
            label = stem.split("_", 1)[-1] if "_" in stem else stem
            excerpts.append(f"＜{label}＞\n{chunk}")
            total += len(chunk)
            if total >= max_chars:
                break
        if total >= max_chars:
            break

    if not excerpts:
        return ""
    return "## 参考文献（読書データ）抜粋\n\n" + "\n\n---\n\n".join(excerpts[:12])


MATERIAL_PROMPT_TEMPLATE = """\
以下の【教科書分析テキスト】と【参考文献抜粋】をもとに、授業の「行間を埋める」文字資料を作成してください。

## 目的
教科書が省いている「なぜ→なぜ→なぜ」の因果連鎖を補い、歴史の流れを生徒が論理的に読める資料を作る。
各ステップで「なぜその次の状況が生まれたのか」が読んで分かること（原因と結果が論理的に落ちない記述）。

## 作成方針
- **因果関係・背景・結果のみ**を扱う（年号の列挙・人物の単純紹介は不要）
- 各ステップには「その行動を取った理由」や「その状況が生まれた構造的背景」を盛り込む
- 1ステップあたり 30〜60 字程度を目安にする（論理的に必要な場合はそれ以上でも可）
- 特に重要な経緯のみ **2〜4 テーマ** を選ぶ（すべてを網羅しない）
- 1テーマにつき矢印チェーンは **3〜6 ステップ**
- 【参考文献抜粋】の内容を積極的に活用し、教科書が省いた構造的説明を補う
- 参考文献で補えない部分は、山川出版社・岩波書店・大学出版局等の信頼できる研究成果に基づいて補う

## 出力形式（厳守）
【テーマ名（15字以内）】
ステップA → ステップB → ステップC → ステップD

テーマ名と矢印チェーンのペアのみ出力。説明文・番号・余計な空行は不要。

## 教科書分析テキスト
{analysis}

{book_section}
"""


def generate_material(analysis: str, session_name: str = "") -> str:
    book_section = load_book_excerpts(session_name) if session_name else ""
    prompt = MATERIAL_PROMPT_TEMPLATE.format(analysis=analysis, book_section=book_section)
    response = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}]
    )
    return response.content[0].text


def save_material_docx(raw_text: str, output_path: Path, session_name: str):
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # タイトル
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_line_spacing(p)
    run = p.add_run(f"歴史の流れ資料　{session_name}")
    _set_font(run, "游ゴシック", 16, bold=True)
    run.font.color.rgb = RGBColor(0x0D, 0x21, 0x37)

    for line in raw_text.splitlines():
        line = line.strip()
        if not line:
            continue

        if line.startswith("【") and "】" in line:
            # テーマ見出し：左ボーダー付き
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after  = Pt(4)
            _set_line_spacing(p)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "12")
            left.set(qn("w:space"), "8")
            left.set(qn("w:color"), "0D2137")
            pBdr.append(left)
            pPr.append(pBdr)
            run = p.add_run(line)
            _set_font(run, "游ゴシック", 13, bold=True)
            run.font.color.rgb = RGBColor(0x0D, 0x21, 0x37)
        else:
            # 矢印チェーン：→ を赤で強調
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(4)
            _set_line_spacing(p, 1.6)
            parts = line.split("→")
            for i, part in enumerate(parts):
                part = part.strip()
                if i > 0:
                    arrow = p.add_run("  →  ")
                    _set_font(arrow, "游ゴシック", 12, bold=True)
                    arrow.font.color.rgb = RGBColor(0x8B, 0x20, 0x20)
                run = p.add_run(part)
                _set_font(run, "游明朝", 12)

    doc.save(output_path)
    sprint(f"  → 文字資料.docx 保存")


# ================================================================
# メイン処理
# ================================================================

OUTPUT_FILES = [
    "一問一答カード.docx",
    "テスト問題.docx",
    "文字資料.docx",
]


def is_done(session_dir: Path) -> bool:
    out_dir = OUTPUT_DIR / session_dir.name
    return all((out_dir / f).exists() for f in OUTPUT_FILES)


def process_session(session_dir: Path, force: bool = False):
    session_name = session_dir.name
    out_dir = OUTPUT_DIR / session_name
    out_dir.mkdir(parents=True, exist_ok=True)

    if not force and is_done(session_dir):
        sprint(f"\n[スキップ] {session_name}  （出力済み。再生成するには --force を付けて実行）")
        return

    sprint(f"\n{'='*60}")
    sprint(f"  [{session_name}] 処理開始")
    sprint(f"{'='*60}")

    # 画像読み込み
    sprint("\n[画像読み込み]")
    image_blocks = load_images(session_dir)
    if not image_blocks:
        sprint("  [スキップ] 画像が見つかりません（.jpg/.jpeg/.png/.webp）")
        return

    # Phase 1: 共有分析（キャッシュあれば再利用）
    cache_path = out_dir / "analysis.txt"
    if cache_path.exists():
        sprint("\n[Phase 1] キャッシュ使用（analysis.txt）")
        analysis = cache_path.read_text(encoding="utf-8")
    else:
        sprint("\n[Phase 1] 教科書分析中... （全ツール共用・約 30 秒）")
        try:
            analysis = analyze_images(image_blocks)
            cache_path.write_text(analysis, encoding="utf-8")
            sprint("  → analysis.txt に保存（次回は再利用）")
        except Exception as e:
            sprint(f"  [エラー] 分析失敗: {e}")
            return

    errors = []

    # Phase 2a: 一問一答
    sprint("\n[Phase 2a] 一問一答カード生成中...")
    try:
        cards = generate_anki_cards(analysis)
        sprint(f"  → {len(cards)}枚生成")
        save_anki_docx(cards, out_dir / "一問一答カード.docx", session_name)
    except Exception as e:
        sprint(f"  [エラー] 一問一答失敗: {e}")
        errors.append(f"一問一答: {e}")

    # Phase 2b: テスト問題
    sprint("\n[Phase 2b] テスト問題生成中... （約 30 秒）")
    try:
        test_raw = generate_test_questions(analysis)
        save_test_docx(test_raw, out_dir / "テスト問題.docx", session_name)
    except Exception as e:
        sprint(f"  [エラー] テスト問題失敗: {e}")
        errors.append(f"テスト問題: {e}")

    # Phase 2c: 文字資料
    sprint("\n[Phase 2c] 文字資料生成中（書籍データ参照）...")
    try:
        material_raw = generate_material(analysis, session_name)
        save_material_docx(material_raw, out_dir / "文字資料.docx", session_name)
    except Exception as e:
        sprint(f"  [エラー] 文字資料失敗: {e}")
        errors.append(f"文字資料: {e}")

    sprint(f"\n{'='*60}")
    sprint(f"  完了！ → output/{session_name}/")
    sprint(f"    一問一答カード.docx")
    sprint(f"    テスト問題.docx")
    sprint(f"    文字資料.docx")
    sprint(f"    analysis.txt（分析キャッシュ）")
    if errors:
        sprint(f"\n  [警告] 一部エラーあり:")
        for err in errors:
            sprint(f"    - {err}")
    sprint(f"{'='*60}")


def main():
    sprint("=" * 60)
    sprint("  授業自動生成ツール")
    sprint("  一問一答カード + テスト問題 + 文字資料")
    sprint("=" * 60)

    if not os.getenv("ANTHROPIC_API_KEY"):
        sprint("\n[エラー] ANTHROPIC_API_KEY が設定されていません。")
        sprint("  .env ファイルに ANTHROPIC_API_KEY=sk-ant-... を記入してください。")
        sys.exit(1)

    INPUT_DIR.mkdir(exist_ok=True)
    OUTPUT_DIR.mkdir(exist_ok=True)

    session_dirs = [d for d in sorted(INPUT_DIR.iterdir()) if d.is_dir()]
    if not session_dirs:
        sprint("\n[注意] input/ に授業回フォルダがありません。")
        sprint("  例: input/01_産業革命/ フォルダを作成し、教科書画像（4枚程度）を入れてください。")
        sprint("\nフォルダ構造:")
        sprint("  input/")
        sprint("    01_産業革命/")
        sprint("      p001.jpg")
        sprint("      p002.jpg")
        sprint("      p003.jpg")
        sprint("      p004.jpg")
        return

    force = "--force" in sys.argv

    sprint(f"\n{len(session_dirs)} 件の授業回フォルダを検出:")
    for d in session_dirs:
        imgs = [p for p in d.iterdir() if p.is_file() and _is_image(p)]
        status = "✓ 出力済み" if is_done(d) else "→ 未処理"
        sprint(f"  {status}  {d.name}/  ({len(imgs)} 枚)")

    if force:
        sprint("\n[--force] 出力済みフォルダも再生成します")

    for session_dir in session_dirs:
        process_session(session_dir, force=force)

    sprint("\n全処理完了！")


if __name__ == "__main__":
    main()
